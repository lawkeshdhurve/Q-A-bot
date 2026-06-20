"""
generate_samples.py
Generates 3 realistic sample documents for testing the RAG pipeline:
  1. business_doc.pdf   - Fictional company annual report
  2. science_paper.pdf   - Fictional research paper on battery technology
  3. factsheet.docx      - Fictional product factsheet

These are NOT real documents - all data, companies, and findings are
fictional and created purely to give the RAG pipeline something
realistic to ingest, chunk, embed, and retrieve from.
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, Inches
import os

OUT_DIR = "data"
os.makedirs(OUT_DIR, exist_ok=True)

styles = getSampleStyleSheet()
h1 = ParagraphStyle("h1", parent=styles["Heading1"], spaceAfter=14)
h2 = ParagraphStyle("h2", parent=styles["Heading2"], spaceAfter=10)
body = ParagraphStyle("body", parent=styles["BodyText"], spaceAfter=10, leading=15)


# ---------------------------------------------------------------------------
# 1. business_doc.pdf - Annual Report (fictional company: Solara Dynamics)
# ---------------------------------------------------------------------------
def build_business_doc():
    path = os.path.join(OUT_DIR, "business_doc.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter,
                             topMargin=0.8*inch, bottomMargin=0.8*inch)
    elems = []

    elems.append(Paragraph("Solara Dynamics Inc.", h1))
    elems.append(Paragraph("Annual Report — Fiscal Year 2026", h2))
    elems.append(Spacer(1, 12))

    elems.append(Paragraph("Executive Summary", h2))
    elems.append(Paragraph(
        "Solara Dynamics Inc. is a renewable energy technology company headquartered in "
        "Pune, Maharashtra, specializing in solar microinverters and battery storage "
        "systems for residential and commercial use. Fiscal Year 2026 marked a period of "
        "significant growth for the company, driven by expansion into Southeast Asian "
        "markets and the launch of the SolFlex 3000 microinverter line.", body))
    elems.append(Paragraph(
        "The company's net revenue for FY2026 grew by 14% year-over-year, reaching "
        "&#8377;842 crore, compared to &#8377;738 crore in FY2025. This growth was "
        "primarily attributed to a 22% increase in unit sales of the SolFlex product "
        "line and improved gross margins following the renegotiation of supplier "
        "contracts in Q2.", body))
    elems.append(PageBreak())

    elems.append(Paragraph("Financial Highlights", h2))
    elems.append(Paragraph(
        "The board of directors approved a dividend payout ratio of 18% for FY2026, "
        "up from 15% in the prior year. Operating expenses increased modestly by 6%, "
        "primarily due to increased R&amp;D investment in next-generation battery "
        "chemistry. The company's debt-to-equity ratio improved to 0.42, down from "
        "0.51 in FY2025, reflecting stronger balance sheet management.", body))

    table_data = [
        ["Metric", "FY2025", "FY2026", "YoY Change"],
        ["Net Revenue (₹ crore)", "738", "842", "+14%"],
        ["Gross Margin", "31.2%", "34.8%", "+3.6 pp"],
        ["Operating Expenses (₹ crore)", "212", "225", "+6%"],
        ["Net Profit (₹ crore)", "94", "118", "+25.5%"],
        ["R&D Spend (₹ crore)", "48", "61", "+27%"],
    ]
    t = Table(table_data, colWidths=[2.4*inch, 1.3*inch, 1.3*inch, 1.3*inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E5E8C")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F4F8")]),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    elems.append(t)
    elems.append(Spacer(1, 16))

    elems.append(Paragraph(
        "The company's flagship product, the SolFlex 3000 microinverter, achieved a "
        "peak conversion efficiency of 98.6%, positioning it competitively against "
        "established players in the European and North American markets.", body))
    elems.append(PageBreak())

    elems.append(Paragraph("Regional Expansion", h2))
    elems.append(Paragraph(
        "In Q3 FY2026, Solara Dynamics opened a new manufacturing facility in Hai "
        "Phong, Vietnam, with an annual production capacity of 2.4 million inverter "
        "units. This facility is expected to reduce logistics costs for Southeast "
        "Asian distribution by an estimated 19% starting FY2027.", body))
    elems.append(Paragraph(
        "The company also signed a strategic distribution agreement with GreenGrid "
        "Solutions, a Jakarta-based energy distributor, to serve the Indonesian "
        "residential solar market. This partnership is projected to contribute "
        "approximately ₹40 crore in incremental annual revenue beginning in FY2027.",
        body))
    elems.append(Spacer(1, 10))

    elems.append(Paragraph("Risk Factors", h2))
    elems.append(Paragraph(
        "Management has identified raw material price volatility, particularly for "
        "silicon wafers and lithium carbonate, as the most significant risk to gross "
        "margin stability in FY2027. The company has partially hedged exposure "
        "through fixed-price contracts covering approximately 60% of projected "
        "silicon demand through Q2 FY2027.", body))
    elems.append(Paragraph(
        "Currency fluctuation also poses a risk given the company's growing exposure "
        "to Vietnamese dong and Indonesian rupiah denominated contracts. The finance "
        "team has implemented forward contracts to hedge approximately 70% of "
        "projected foreign currency receivables for the next two fiscal quarters.",
        body))
    elems.append(PageBreak())

    elems.append(Paragraph("Outlook for FY2027", h2))
    elems.append(Paragraph(
        "Management projects FY2027 net revenue growth of 18-22%, driven by full-year "
        "contribution from the Vietnam facility and continued expansion of the "
        "SolFlex product line, including the planned Q1 FY2027 launch of the SolFlex "
        "5000, a higher-capacity unit targeting commercial rooftop installations.",
        body))
    elems.append(Paragraph(
        "The board has approved a capital expenditure budget of ₹95 crore for FY2027, "
        "primarily allocated toward automation upgrades at the Pune facility and "
        "continued R&amp;D investment in solid-state battery research.", body))

    doc.build(elems)
    print(f"Created {path}")


# ---------------------------------------------------------------------------
# 2. science_paper.pdf - Research paper (fictional: solid-state batteries)
# ---------------------------------------------------------------------------
def build_science_paper():
    path = os.path.join(OUT_DIR, "science_paper.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter,
                             topMargin=0.8*inch, bottomMargin=0.8*inch)
    elems = []

    elems.append(Paragraph(
        "Lithium-Lanthanum-Zirconate Solid Electrolytes for High-Density "
        "Energy Storage: A Comparative Study", h1))
    elems.append(Paragraph(
        "R. Krishnan, A. Bhattacharya, and M. Okafor &mdash; Department of "
        "Materials Engineering, Indian Institute of Science Education", body))
    elems.append(Spacer(1, 10))

    elems.append(Paragraph("Abstract", h2))
    elems.append(Paragraph(
        "This study investigates the ionic conductivity and electrochemical "
        "stability of garnet-type lithium-lanthanum-zirconate (LLZO) solid "
        "electrolytes as a candidate material for next-generation solid-state "
        "lithium batteries. We report a room-temperature ionic conductivity of "
        "1.2 &times; 10&minus;3 S/cm for aluminum-doped LLZO samples sintered at "
        "1230&deg;C, representing a 35% improvement over previously reported "
        "undoped LLZO formulations.", body))
    elems.append(PageBreak())

    elems.append(Paragraph("1. Introduction", h2))
    elems.append(Paragraph(
        "Conventional lithium-ion batteries rely on liquid organic electrolytes, "
        "which pose safety risks including flammability and electrolyte leakage. "
        "Solid-state electrolytes have emerged as a promising alternative, "
        "offering improved thermal stability and the potential for higher energy "
        "density through compatibility with lithium metal anodes.", body))
    elems.append(Paragraph(
        "Among solid electrolyte candidates, garnet-type LLZO has attracted "
        "considerable attention due to its high ionic conductivity, wide "
        "electrochemical stability window, and chemical compatibility with "
        "lithium metal. However, challenges remain in achieving consistent "
        "grain boundary conductivity and minimizing interfacial resistance "
        "with electrode materials.", body))

    elems.append(Paragraph("2. Methodology", h2))
    elems.append(Paragraph(
        "Samples of aluminum-doped LLZO (Li6.25Al0.25La3Zr2O12) were synthesized "
        "via solid-state reaction using high-purity precursor powders. The "
        "powders were calcined at 950&deg;C for 6 hours, ball-milled, and "
        "subsequently sintered at temperatures ranging from 1150&deg;C to "
        "1280&deg;C for 8 hours under an argon atmosphere to minimize lithium "
        "volatilization.", body))
    elems.append(Paragraph(
        "Ionic conductivity was measured using electrochemical impedance "
        "spectroscopy (EIS) across a frequency range of 1 Hz to 1 MHz at "
        "temperatures from 25&deg;C to 100&deg;C. Sample density was determined "
        "using the Archimedes method, and phase purity was confirmed via X-ray "
        "diffraction (XRD).", body))
    elems.append(PageBreak())

    elems.append(Paragraph("3. Results", h2))
    elems.append(Paragraph(
        "Samples sintered at 1230&deg;C exhibited the highest relative density "
        "at 96.4%, correlating with the highest measured ionic conductivity of "
        "1.2 &times; 10&minus;3 S/cm at room temperature. Samples sintered below "
        "1180&deg;C showed significantly reduced density (below 88%) and "
        "correspondingly lower conductivity, attributed to incomplete "
        "densification and increased grain boundary resistance.", body))
    elems.append(Paragraph(
        "The activation energy for lithium-ion migration, calculated from "
        "Arrhenius plots, was found to be 0.31 eV for the optimally sintered "
        "samples, consistent with values reported for high-performance garnet "
        "electrolytes in prior literature. Cyclic voltammetry confirmed an "
        "electrochemical stability window extending to approximately 6 V "
        "relative to Li/Li+, indicating compatibility with high-voltage cathode "
        "materials.", body))

    elems.append(Paragraph("4. Discussion", h2))
    elems.append(Paragraph(
        "The observed improvement in ionic conductivity at the 1230&deg;C "
        "sintering condition is attributed to optimal grain growth and reduced "
        "porosity, which together minimize obstruction of lithium-ion transport "
        "pathways across grain boundaries. These findings suggest that sintering "
        "temperature is a critical and highly sensitive parameter in optimizing "
        "LLZO performance for practical solid-state battery applications.", body))
    elems.append(PageBreak())

    elems.append(Paragraph("5. Conclusion", h2))
    elems.append(Paragraph(
        "Aluminum-doped LLZO sintered at 1230&deg;C demonstrates ionic "
        "conductivity and electrochemical stability sufficient for "
        "consideration in next-generation solid-state lithium battery designs. "
        "Future work should focus on reducing interfacial resistance between "
        "the LLZO electrolyte and lithium metal anodes, as well as scaling "
        "synthesis methods for industrial production.", body))

    doc.build(elems)
    print(f"Created {path}")


# ---------------------------------------------------------------------------
# 3. factsheet.docx - Product factsheet (fictional: SolFlex 3000 inverter)
# ---------------------------------------------------------------------------
def build_factsheet():
    path = os.path.join(OUT_DIR, "factsheet.docx")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("SolFlex 3000 Microinverter — Product Factsheet", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "The SolFlex 3000 is a grid-tied solar microinverter designed for "
        "residential rooftop installations. It converts DC power generated by "
        "individual solar panels into grid-compatible AC power, offering "
        "panel-level monitoring and improved system resilience compared to "
        "traditional string inverter designs."
    )

    doc.add_heading("Technical Specifications", level=2)
    specs = [
        ("Rated Output Power", "300 W"),
        ("Peak Efficiency", "98.6%"),
        ("Input Voltage Range", "16V – 60V DC"),
        ("Output Voltage", "230V AC (configurable for 110V regions)"),
        ("Operating Temperature Range", "-25°C to 60°C"),
        ("Weight", "1.2 kg"),
        ("Warranty Period", "12 years standard, extendable to 20 years"),
        ("Communication Protocol", "Zigbee mesh network with cloud monitoring"),
        ("IP Rating", "IP67 (fully weatherproof)"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Specification"
    hdr[1].text = "Value"
    for spec, val in specs:
        row = table.add_row().cells
        row[0].text = spec
        row[1].text = val

    doc.add_heading("Installation Notes", level=2)
    doc.add_paragraph(
        "Each SolFlex 3000 unit should be mounted directly beneath its "
        "corresponding solar panel using the supplied mounting bracket. "
        "Installers must ensure a minimum clearance of 2 cm between the "
        "inverter housing and the roof surface to allow for adequate passive "
        "cooling airflow."
    )
    doc.add_paragraph(
        "The Zigbee mesh network supports a maximum of 50 microinverters per "
        "monitoring gateway. For installations exceeding 50 panels, an "
        "additional gateway unit must be installed, with overlapping mesh "
        "coverage recommended to maintain communication reliability."
    )

    doc.add_heading("Safety Certifications", level=2)
    doc.add_paragraph(
        "The SolFlex 3000 has been certified to IEC 62109-1 and IEC 62109-2 "
        "safety standards for power converters used in photovoltaic systems. "
        "It also holds BIS certification (IS 16221) for sale within the Indian "
        "market and CE certification for European Union markets."
    )

    doc.add_heading("Warranty and Support", level=2)
    doc.add_paragraph(
        "Solara Dynamics provides a standard 12-year limited warranty covering "
        "manufacturing defects, with an optional extended warranty available up "
        "to 20 years at the time of purchase. Customers experiencing a unit "
        "failure under warranty should contact regional support with the unit's "
        "serial number and installation date for expedited replacement "
        "processing, typically completed within 5 business days."
    )

    doc.save(path)
    print(f"Created {path}")


if __name__ == "__main__":
    build_business_doc()
    build_science_paper()
    build_factsheet()
    print("\nAll sample documents generated in ./data/")
