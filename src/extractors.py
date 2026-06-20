"""
extractors.py
Handles raw text extraction from supported document formats (PDF, DOCX, TXT).

Each extractor function returns a list of dictionaries in a consistent shape:
    {"text": "<page or section text>", "metadata": {"source": "<filename>", "page": <int>}}

Keeping a uniform output shape means downstream chunking logic (chunking.py)
doesn't need to care which file format the text originally came from.
"""

import os
from pypdf import PdfReader
from docx import Document as DocxDocument


def extract_pdf_pages(file_path: str) -> list[dict]:
    """
    Extracts text page-by-page from a PDF, tracking page numbers and file source.

    Why page-by-page: it lets us cite the exact page a fact came from later,
    which is essential for the "Source: file.pdf, Page: N" citations the
    assignment requires.
    """
    extracted_data = []
    file_name = os.path.basename(file_path)

    try:
        reader = PdfReader(file_path)
        for index, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                # Collapse multiple whitespace/newlines into single spaces.
                # Raw PDF text extraction often has irregular line breaks
                # from the original page layout.
                clean_text = " ".join(text.split())
                extracted_data.append({
                    "text": clean_text,
                    "metadata": {
                        "source": file_name,
                        "page": index + 1  # 1-indexed for human-readable citations
                    }
                })
    except Exception as e:
        print(f"  [Error] Could not read PDF {file_name}: {e}")

    return extracted_data


def extract_docx_pages(file_path: str) -> list[dict]:
    """
    Extracts text from a DOCX file.

    DOCX files don't have a native "page" concept the way PDFs do (page
    breaks depend on rendering/font/viewport), so we group all paragraph
    text into a single logical unit per file, with page=1. If a single
    DOCX is very large, the chunking step below still slices it into
    multiple smaller chunks - we just won't have true page numbers as
    citation anchors, only the source filename.
    """
    extracted_data = []
    file_name = os.path.basename(file_path)

    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Also pull text out of any tables, since python-docx doesn't
        # include table cell text in doc.paragraphs.
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text.append(row_text)

        full_text = "\n".join(paragraphs + table_text)
        full_text = " ".join(full_text.split())  # normalize whitespace

        if full_text:
            extracted_data.append({
                "text": full_text,
                "metadata": {
                    "source": file_name,
                    "page": 1
                }
            })
    except Exception as e:
        print(f"  [Error] Could not read DOCX {file_name}: {e}")

    return extracted_data


def extract_txt_pages(file_path: str) -> list[dict]:
    """Extracts text from a plain .txt file (bonus format support)."""
    extracted_data = []
    file_name = os.path.basename(file_path)

    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        clean_text = " ".join(content.split())
        if clean_text:
            extracted_data.append({
                "text": clean_text,
                "metadata": {"source": file_name, "page": 1}
            })
    except Exception as e:
        print(f"  [Error] Could not read TXT {file_name}: {e}")

    return extracted_data


def extract_document(file_path: str) -> list[dict]:
    """
    Dispatcher: routes a file to the correct extractor based on its extension.
    Returns an empty list (with a warning) for unsupported file types,
    rather than crashing the whole ingestion run over one bad file.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_pdf_pages(file_path)
    elif ext == ".docx":
        return extract_docx_pages(file_path)
    elif ext == ".txt":
        return extract_txt_pages(file_path)
    else:
        print(f"  [Skipped] Unsupported file type: {file_path}")
        return []
